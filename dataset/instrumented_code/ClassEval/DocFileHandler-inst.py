import inspect
import json
import os
from datetime import datetime

def custom_serializer(obj):
    if isinstance(obj, datetime):
        return obj.isoformat()
    return str(obj)


def recursive_object_seralizer(obj, visited):
    seralized_dict = {}
    keys = list(obj.__dict__)
    for k in keys:
        if id(obj.__dict__[k]) in visited:
            seralized_dict[k] = "<RECURSIVE {}>".format(obj.__dict__[k])
            continue
        if isinstance(obj.__dict__[k], (float, int, str, bool, type(None))):
            seralized_dict[k] = obj.__dict__[k]
        elif isinstance(obj.__dict__[k], tuple):
            ## handle tuple
            seralized_dict[k] = obj.__dict__[k]
        elif isinstance(obj.__dict__[k], set):
            ## handle set
            seralized_dict[k] = obj.__dict__[k]
        elif isinstance(obj.__dict__[k], list):
            ## handle list
            seralized_dict[k] = obj.__dict__[k]
        elif hasattr(obj.__dict__[k], '__dict__'):
            ## handle object
            visited.append(id(obj.__dict__[k]))
            seralized_dict[k] = obj.__dict__[k]
        elif isinstance(obj.__dict__[k], dict):
            visited.append(id(obj.__dict__[k]))
            seralized_dict[k] = obj.__dict__[k]
        elif callable(obj.__dict__[k]):
            ## handle function
            if hasattr(obj.__dict__[k], '__name__'):
                seralized_dict[k] = "<function {}>".format(obj.__dict__[k].__name__)
        else:
            seralized_dict[k] = str(obj.__dict__[k])
    return seralized_dict

def inspect_code(func):
   def wrapper(*args, **kwargs):
       visited = []
       json_base = "/home/changshu/ClassEval/data/benchmark_solution_code/input-output/"
       if not os.path.exists(json_base):
           os.mkdir(json_base)
       jsonl_path = json_base + "/DocFileHandler.jsonl"
       para_dict = {"name": func.__name__}
       args_names = inspect.getfullargspec(func).args
       if len(args) > 0 and hasattr(args[0], '__dict__') and args_names[0] == 'self':
           ## 'self'
           self_args = args[0]
           para_dict['self'] = recursive_object_seralizer(self_args, [id(self_args)])
       else:
           para_dict['self'] = {}
       if len(args) > 0 :
           if args_names[0] == 'self':
               other_args = {}
               for m,n in zip(args_names[1:], args[1:]):
                   other_args[m] = n
           else:
               other_args = {}
               for m,n in zip(args_names, args):
                   other_args[m] = n
           
           para_dict['args'] = other_args
       else:
           para_dict['args'] = {}
       if kwargs:
           para_dict['kwargs'] = kwargs
       else:
           para_dict['kwargs'] = {}
          
       result = func(*args, **kwargs)
       para_dict["return"] = result
       with open(jsonl_path, 'a') as f:
           f.write(json.dumps(para_dict, default=custom_serializer) + "\n")
       return result
   return wrapper


'''
# This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
'''

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path

    @inspect_code
    def read_text(self):
        doc = Document(self.file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text)

    @inspect_code
    def write_text(self, content, font_size=12, alignment='left'):
        try:
            doc = Document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(content)
            font = run.font
            font.size = Pt(font_size)
            alignment_value = self._get_alignment_value(alignment)
            paragraph.alignment = alignment_value
            doc.save(self.file_path)
            return True
        except:
            return False

    @inspect_code
    def add_heading(self, heading, level=1):
        try:
            doc = Document(self.file_path)
            doc.add_heading(heading, level)
            doc.save(self.file_path)
            return True
        except:
            return False

    @inspect_code
    def add_table(self, data):
        try:
            doc = Document(self.file_path)
            table = doc.add_table(rows=len(data), cols=len(data[0]))
            for i, row in enumerate(data):
                for j, cell_value in enumerate(row):
                    table.cell(i, j).text = str(cell_value)
            doc.save(self.file_path)
            return True
        except:
            return False

    @inspect_code
    def _get_alignment_value(self, alignment):
        alignment_options = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_options.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)



import unittest
import os


class DocFileHandlerTestReadText(unittest.TestCase):
    def test_read_text_1(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("Initial content")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = "Initial content"
        self.assertEqual(text_content, expected_content)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_read_text_2(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("111")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = "111"
        self.assertEqual(text_content, expected_content)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_read_text_3(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("aaa")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = "aaa"
        self.assertEqual(text_content, expected_content)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_read_text_4(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("aaa\nbbb")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = "aaa\nbbb"
        self.assertEqual(text_content, expected_content)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_read_text_5(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = ""
        self.assertEqual(text_content, expected_content)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)


class DocFileHandlerTestWriteText(unittest.TestCase):
    def setUp(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("Initial content")
        doc.save(self.file_path)

    def tearDown(self):
        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_write_text_1(self):
        new_content = "New content 1"
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)

    def test_write_text_2(self):
        new_content = "New content 2"
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)

    def test_write_text_3(self):
        new_content = "New content 3"
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)

    def test_write_text_4(self):
        new_content = "New content 4"
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)

    def test_write_text_5(self):
        new_content = ""
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)


class DocFileHandlerTestAddHeading(unittest.TestCase):
    def setUp(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("Initial content")
        doc.save(self.file_path)

    def tearDown(self):
        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_add_heading_1(self):
        heading = "Test Heading 1"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)

    def test_add_heading_2(self):
        heading = "Test Heading 2"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)

    def test_add_heading_3(self):
        heading = "Test Heading 3"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)

    def test_add_heading_4(self):
        heading = "Test Heading 4"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)

    def test_add_heading_5(self):
        heading = "Test Heading 5"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)


class DocFileHandlerTestAddTable(unittest.TestCase):
    def setUp(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("Initial content")
        doc.save(self.file_path)

    def tearDown(self):
        if os.path.exists(self.file_path):
            os.remove(self.file_path)

    def test_add_table_1(self):
        data = [['Name', 'Age']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 1)
        self.assertEqual(len(table.columns), 2)

    def test_add_table_2(self):
        data = [['Name', 'Age'], ['John', '25']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 2)
        self.assertEqual(len(table.columns), 2)
        self.assertEqual(table.cell(1, 0).text, 'John')

    def test_add_table_3(self):
        data = [['Name', 'Age'], ['John', '25'], ['Emma', '30']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 3)
        self.assertEqual(len(table.columns), 2)
        self.assertEqual(table.cell(1, 0).text, 'John')
        self.assertEqual(table.cell(2, 1).text, '30')

    def test_add_table_4(self):
        data = [['Name', 'Age'], ['aaa', '25'], ['Emma', '30']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 3)
        self.assertEqual(len(table.columns), 2)
        self.assertEqual(table.cell(1, 0).text, 'aaa')
        self.assertEqual(table.cell(2, 1).text, '30')

    def test_add_table_5(self):
        data = [['Name', 'Age'], ['John', '25'], ['Emma', '90']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 3)
        self.assertEqual(len(table.columns), 2)
        self.assertEqual(table.cell(1, 0).text, 'John')
        self.assertEqual(table.cell(2, 1).text, '90')


class DocFileHandlerTest(unittest.TestCase):
    def test_DocFileHandler(self):
        self.file_path = "test_example.docx"
        self.handler = DocFileHandler(self.file_path)
        doc = Document()
        doc.add_paragraph("Initial content")
        doc.save(self.file_path)

        text_content = self.handler.read_text()
        expected_content = "Initial content"
        self.assertEqual(text_content, expected_content)

        new_content = "New content 1"
        self.handler.write_text(new_content)
        text_content = self.handler.read_text()
        self.assertEqual(text_content, new_content)

        heading = "Test Heading 1"
        self.handler.add_heading(heading)
        doc = Document(self.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        self.assertIn(heading, headings)

        data = [['Name', 'Age']]
        self.handler.add_table(data)
        doc = Document(self.file_path)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 1)
        self.assertEqual(len(table.columns), 2)

        if os.path.exists(self.file_path):
            os.remove(self.file_path)

