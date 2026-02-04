
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocFileHandler():

    def __init__(self, file_path):
        self.file_path = file_path

    def add_table(self, data):
        try:
            doc = Document(self.file_path)
            table = doc.add_table(rows=len(data), cols=len(data[0]))
            for (i, row) in enumerate(data):
                for (j, cell_value) in enumerate(row):
                    table.cell(i, j).text = str(cell_value)
            doc.save(self.file_path)
            return True
        except:
            return False
